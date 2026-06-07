"""
exporters/word_exporter.py
python-docx 기반 Word(.docx) 내보내기 — Node.js 불필요
"""
from __future__ import annotations
import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 색상 상수 ─────────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x0D, 0x1B, 0x4B)
C_BLUE   = RGBColor(0x15, 0x65, 0xC0)
C_LGRAY  = RGBColor(0x54, 0x6E, 0x7A)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_RED    = RGBColor(0xC6, 0x28, 0x28)
C_DBLUE  = RGBColor(0x15, 0x65, 0xC0)
C_BLACK  = RGBColor(0x1A, 0x1A, 0x2E)
C_GOLD   = RGBColor(0xFF, 0xD5, 0x4F)


def _set_cell_bg(cell, hex_color: str):
    """표 셀 배경색 지정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = C_NAVY if level == 1 else C_BLUE
    return p


def _add_body(doc: Document, text: str, bold: bool = False, color: RGBColor = None):
    if not text.strip():
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.replace("**", "").replace("##", "").strip())
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or C_BLACK
    return p


def export_word(
    report_text: str,
    stocks: list[dict],
    persona_label: str,
    output_path: str,
) -> str:
    today = date.today().isoformat()
    doc = Document()

    # ── 페이지 마진 ───────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # ── 표지 ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    tr = title_p.add_run("📊 주식 AI 투자 브리핑 리포트")
    tr.bold = True; tr.font.size = Pt(22); tr.font.color.rgb = C_NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run(f"기준일: {today}  |  분석 관점: {persona_label}")
    sr.font.size = Pt(11); sr.font.color.rgb = C_LGRAY

    doc.add_paragraph()

    # ── 종목 현황 표 ──────────────────────────────────────────────────────────
    _add_heading(doc, "📈 종목별 주가 현황", level=1)

    headers = ["종목명", "코드/티커", "종가", "등락률", "시장"]
    col_widths = [Cm(4.5), Cm(2.5), Cm(2.8), Cm(2.5), Cm(1.5)]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # 헤더 행
    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        _set_cell_bg(cell, "1565C0")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = C_WHITE

    # 데이터 행
    for row_idx, s in enumerate(stocks[:10]):
        row = tbl.add_row()
        name   = s.get("name", s["ticker"])
        ticker = s["ticker"]
        close  = s.get("close")
        change = s.get("change_rate")
        market = s.get("market", "")
        currency = "원" if market == "KR" else "$"
        close_str  = f"{close:,.0f}{currency}" if close else "N/A"
        change_str = f"{change:+.2f}%" if change is not None else "N/A"
        change_color = C_RED if (change or 0) > 0 else C_DBLUE if (change or 0) < 0 else C_LGRAY

        values = [name, ticker, close_str, change_str, market]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                  WD_ALIGN_PARAGRAPH.CENTER]
        colors = [C_BLACK, C_BLACK, C_BLACK, change_color, C_BLACK]

        bg = "F5F9FF" if row_idx % 2 == 0 else "FFFFFF"
        for i, (val, align, color) in enumerate(zip(values, aligns, colors)):
            cell = row.cells[i]
            cell.width = col_widths[i]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = align
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = color
            if i == 3:  # 등락률 굵게
                run.bold = True

    doc.add_paragraph()

    # ── 포트폴리오 종합 전략 ──────────────────────────────────────────────────
    # (report_text에서 섹션 파싱)
    lines = report_text.split("\n")
    current_section = []
    in_portfolio = False
    in_stock = False

    for line in lines:
        stripped = line.strip()

        # 구분선 처리
        if stripped.startswith("=") and len(stripped) > 20:
            if current_section:
                current_section = []
            continue
        if stripped.startswith("-") and len(stripped) > 20:
            continue

        # 섹션 헤더
        if "포트폴리오 종합 전략" in stripped:
            _add_heading(doc, "📋 포트폴리오 종합 전략", level=1)
            in_portfolio = True
            in_stock = False
            continue

        if stripped.startswith("[ ") and stripped.endswith(" ]"):
            # 종목 개별 섹션
            section_title = stripped.strip("[]").strip()
            _add_heading(doc, f"📌 {section_title}", level=1)
            in_portfolio = False
            in_stock = True
            continue

        # 본문 처리
        if not stripped:
            doc.add_paragraph()
            continue

        # 굵은 항목 (숫자. 또는 ** 또는 ## 시작)
        is_heading_line = (
            stripped.startswith("**") or
            stripped.startswith("##") or
            (len(stripped) > 2 and stripped[1] == "." and stripped[0].isdigit())
        )
        clean = stripped.replace("**", "").replace("##", "").strip()
        if clean:
            _add_body(doc, clean, bold=is_heading_line,
                      color=C_BLUE if is_heading_line else C_BLACK)

    # ── 푸터 면책 ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_r = disc_p.add_run(
        f"본 리포트는 AI 분석 기반이며 투자 결정은 본인 책임입니다.  |  {today}"
    )
    disc_r.font.size = Pt(9)
    disc_r.font.color.rgb = C_LGRAY
    disc_r.italic = True

    # ── 저장 ─────────────────────────────────────────────────────────────────
    doc.save(output_path)
    return output_path


def export_word_bytes(report_text: str, stocks: list[dict], persona_label: str) -> bytes:
    """바이트로 반환 (Streamlit download_button용)."""
    import tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp = f.name
    try:
        export_word(report_text, stocks, persona_label, tmp)
        return open(tmp, "rb").read()
    finally:
        os.unlink(tmp)
